import os
import json
import logging
import pandas as pd
from datetime import datetime
from typing import List, Optional
from sqlalchemy import create_engine, Column, Integer, String, Float, Text, DateTime, desc
from sqlalchemy.orm import sessionmaker, declarative_base
from config.settings import settings
from models.schemas import EscalationTask, GeneratedEmail, AuditRecord

logger = logging.getLogger(__name__)

# SQLAlchemy Setup
Base = declarative_base()

class EmailAudit(Base):
    __tablename__ = 'email_audit'
    id = Column(Integer, primary_key=True, autoincrement=True)
    invoice_no = Column(String(50), nullable=False)
    client_name = Column(String(255), nullable=False)
    client_email = Column(String(255), nullable=False)
    amount_due = Column(Float, nullable=False)
    currency = Column(String(10), default='INR')
    days_overdue = Column(Integer, nullable=False)
    stage = Column(Integer, nullable=False)
    tone_used = Column(String(100))
    subject = Column(String(255))
    body = Column(Text)
    send_status = Column(String(50), nullable=False)
    timestamp = Column(String(100), nullable=False)
    error_message = Column(Text)
    run_id = Column(String(100))

# Database Connection
# Fallback to local SQLite if DATABASE_URL is not set
db_url = os.getenv("DATABASE_URL")
if not db_url:
    db_url = f"sqlite:///{settings.AUDIT_DB_PATH}"
elif db_url.startswith("postgres://"):
    # Fix for newer SQLAlchemy versions requiring 'postgresql://'
    db_url = db_url.replace("postgres://", "postgresql://", 1)

engine = create_engine(db_url)
Session = sessionmaker(bind=engine)

def init_db():
    Base.metadata.create_all(engine)

def log_event(task: EscalationTask, email: Optional[GeneratedEmail] = None, 
              send_status: str = "PENDING", error: Optional[str] = None, run_id: Optional[str] = None):
    
    timestamp = datetime.now().isoformat()
    inv = task.invoice
    
    session = Session()
    try:
        new_record = EmailAudit(
            invoice_no=inv.invoice_no,
            client_name=inv.client_name,
            client_email=inv.client_email,
            amount_due=inv.amount_due,
            currency=inv.currency,
            days_overdue=task.days_overdue,
            stage=task.stage.value,
            tone_used=email.tone_used if email else None,
            subject=email.subject if email else None,
            body=email.body if email else None,
            send_status=send_status,
            timestamp=timestamp,
            error_message=error,
            run_id=run_id
        )
        session.add(new_record)
        session.commit()
    except Exception as e:
        logger.error(f"Database log error: {e}")
        session.rollback()
    finally:
        session.close()
    
    # JSON backup (for local debugging)
    try:
        date_str = datetime.now().strftime("%Y-%m-%d")
        os.makedirs(settings.LOGS_DIR, exist_ok=True)
        json_path = os.path.join(settings.LOGS_DIR, f"audit_{date_str}.json")
        
        audit_entry = {
            "invoice_no": inv.invoice_no,
            "client": inv.client_name,
            "status": send_status,
            "timestamp": timestamp,
            "run_id": run_id
        }
        
        logs = []
        if os.path.exists(json_path):
            with open(json_path, "r") as f:
                logs = json.load(f)
        logs.append(audit_entry)
        with open(json_path, "w") as f:
            json.dump(logs, f, indent=2)
    except:
        pass

def get_all_records() -> List[AuditRecord]:
    session = Session()
    try:
        results = session.query(EmailAudit).order_by(desc(EmailAudit.timestamp)).all()
        records = []
        for r in results:
            # Convert SQLAlchemy object to dict for Pydantic
            r_dict = {c.name: getattr(r, c.name) for c in r.__table__.columns}
            records.append(AuditRecord(**r_dict))
        return records
    finally:
        session.close()

def get_export_data():
    session = Session()
    try:
        results = session.query(EmailAudit).order_by(desc(EmailAudit.timestamp)).all()
        return [{
            "Date": r.timestamp[:19] if r.timestamp else "",
            "Invoice": r.invoice_no,
            "Client": r.client_name,
            "Amount": r.amount_due,
            "Days Overdue": r.days_overdue,
            "Stage": f"Stage {r.stage}" if r.stage < 5 else "Legal",
            "Status": r.send_status,
            "Tone": r.tone_used or "",
            "Subject": r.subject or "",
            "Body": r.body or "",
            "Error Message": r.error_message or "",
            "Run ID": r.run_id or ""
        } for r in results]
    finally:
        session.close()

def export_to_excel(output_path: str):
    data = get_export_data()
    df = pd.DataFrame(data)
    with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Audit Log', index=False)
        workbook = writer.book
        worksheet = writer.sheets['Audit Log']
        
        # Formats
        header_fmt = workbook.add_format({'bold': True, 'bg_color': '#1E293B', 'font_color': 'white', 'border': 1})
        cell_fmt = workbook.add_format({'border': 1})
        money_fmt = workbook.add_format({'num_format': '$#,##0.00', 'border': 1})
        
        # Apply header format
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_fmt)
        
        # Column widths
        worksheet.set_column('A:A', 20, cell_fmt)
        worksheet.set_column('B:B', 15, cell_fmt)
        worksheet.set_column('C:C', 25, cell_fmt)
        worksheet.set_column('D:D', 15, money_fmt)
        worksheet.set_column('E:E', 15, cell_fmt)
        worksheet.set_column('F:F', 12, cell_fmt)
        worksheet.set_column('G:G', 20, cell_fmt)
        worksheet.set_column('H:H', 15, cell_fmt)
        worksheet.set_column('I:I', 40, cell_fmt)
        worksheet.set_column('J:J', 60, cell_fmt)
        worksheet.set_column('K:K', 30, cell_fmt)
        worksheet.set_column('L:L', 36, cell_fmt)

def export_to_pdf(output_path: str):
    try:
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        doc = SimpleDocTemplate(output_path, pagesize=landscape(letter),
                                rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
        elements = []
        styles = getSampleStyleSheet()
        elements.append(Paragraph("Knot AI - Audit Log Report", styles['Title']))
        elements.append(Spacer(1, 12))
        
        data = get_export_data()
        if not data:
            elements.append(Paragraph("No records found.", styles['Normal']))
            doc.build(elements)
            return

        headers = list(data[0].keys())
        
        # Create a custom style for table cells to allow wrapping
        cell_style = ParagraphStyle(
            name='TableCell',
            parent=styles['Normal'],
            fontSize=8,
            leading=10,
            wordWrap='CJK'
        )
        
        header_style = ParagraphStyle(
            name='TableHeader',
            parent=styles['Normal'],
            fontSize=9,
            leading=11,
            textColor=colors.whitesmoke,
            fontName='Helvetica-Bold',
            alignment=1 # Center
        )

        table_data = [[Paragraph(h, header_style) for h in headers]]
        for row in data:
            row_data = []
            for h in headers:
                val = str(row[h]) if row[h] is not None else ""
                row_data.append(Paragraph(val, cell_style))
            table_data.append(row_data)
            
        # Calculate optimal column widths based on available width (11 inches - margins)
        avail_width = landscape(letter)[0] - 60
        # Weights for columns: Date(2), Inv(2), Client(2), Amt(1.5), Days(1), Stage(1.5), Status(2), Tone(1.5), Subj(3), Body(5), Err(3), Run(2)
        weights = [1.5, 1.5, 2.0, 1.0, 1.0, 1.0, 1.5, 1.0, 2.5, 4.0, 2.0, 1.0]
        total_weight = sum(weights)
        col_widths = [(w / total_weight) * avail_width for w in weights]
            
        t = Table(table_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1E293B")),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,0), 8),
            ('TOPPADDING', (0,0), (-1,0), 8),
            ('BACKGROUND', (0,1), (-1,-1), colors.HexColor("#F8FAFC")),
            ('GRID', (0,0), (-1,-1), 1, colors.HexColor("#E2E8F0")),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#F8FAFC")]),
        ]))
        elements.append(t)
        doc.build(elements)
    except Exception as e:
        logger.error(f"PDF Export error: {e}")

def export_to_word(output_path: str):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        
        doc = Document()
        doc.add_heading('Knot AI - Audit Log Report', 0)
        
        data = get_export_data()
        if not data:
            doc.add_paragraph('No records found.')
            doc.save(output_path)
            return
            
        headers = list(data[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        # Add rows
        for item in data:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                row_cells[i].text = str(item[header])
                
        # Adjust simple styling
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        
        doc.save(output_path)
    except Exception as e:
        logger.error(f"Word Export error: {e}")

def clear_all_records():
    """Wipes the database to start fresh."""
    session = Session()
    try:
        session.query(EmailAudit).delete()
        session.commit()
    except Exception as e:
        logger.error(f"Failed to clear db: {e}")
        session.rollback()
    finally:
        session.close()

def resolve_invoice(invoice_no: str):
    """Marks an invoice as manually resolved / paid."""
    timestamp = datetime.now().isoformat()
    session = Session()
    try:
        new_record = EmailAudit(
            invoice_no=invoice_no,
            client_name="RESOLVED",
            client_email="RESOLVED",
            amount_due=0.0,
            currency="USD",
            days_overdue=0,
            stage=0,
            send_status="RESOLVED",
            timestamp=timestamp
        )
        session.add(new_record)
        session.commit()
    except Exception as e:
        logger.error(f"Failed to resolve invoice {invoice_no}: {e}")
        session.rollback()
    finally:
        session.close()
